import io
import os
import re
import time
import hashlib
from dataclasses import dataclass
from typing import List, Tuple, Dict

import numpy as np
import streamlit as st
from openai import OpenAI

# Dependencias opcionales
try:
    from docx import Document
except Exception:
    Document = None

try:
    from PyPDF2 import PdfReader
except Exception:
    PdfReader = None


# ============================================================
# CONFIGURACIÓN GENERAL
# ============================================================
st.set_page_config(
    page_title="Asistente de Atención al Cliente para Guías",
    page_icon="📘",
    layout="wide",
)

APP_TITLE = "📘 Asistente de Atención al Cliente para Guías"
APP_SUBTITLE = "Consultas operativas basadas exclusivamente en el documento vigente."
DEFAULT_MODEL = "gpt-5.1"
DEFAULT_REASONING_EFFORT = "medium"
DEFAULT_EMBEDDING_MODEL = "text-embedding-3-small"
MAX_CHUNK_CHARS = 1700
CHUNK_OVERLAP_CHARS = 220
TOP_K = 10
CONTACTS_TOP_K = 12
REPO_DOC_PATH = "pautasattguias.gpt.docx"


# ============================================================
# MODELOS DE DATOS
# ============================================================
@dataclass
class Chunk:
    chunk_id: str
    heading_path: str
    text: str
    source_name: str


class RepoFile:
    def __init__(self, path: str):
        self.path = path
        self.name = os.path.basename(path)

    def read(self) -> bytes:
        with open(self.path, "rb") as f:
            return f.read()


# ============================================================
# UTILIDADES
# ============================================================
def normalize_whitespace(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = re.sub(r"\r", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def file_sha256(file_bytes: bytes) -> str:
    return hashlib.sha256(file_bytes).hexdigest()


def cosine_similarity_matrix(query_vec: np.ndarray, matrix: np.ndarray) -> np.ndarray:
    query_norm = np.linalg.norm(query_vec)
    matrix_norm = np.linalg.norm(matrix, axis=1)
    denom = (matrix_norm * query_norm) + 1e-12
    return (matrix @ query_vec) / denom


def safe_get_api_key() -> str:
    try:
        return st.secrets.get("OPENAI_API_KEY", "")
    except Exception:
        return ""


def normalize_for_match(text: str) -> str:
    return normalize_whitespace(text).lower()


# ============================================================
# DETECCIÓN DE INTENCIÓN
# ============================================================
def detect_intent(question: str) -> str:
    q = question.lower()

    critical_terms = [
        "accidente", "salud", "médic", "medic", "hospital", "ambul", "emergenc", "urgenc",
        "seguro", "asistencia", "enfermo", "enferma", "doente", "saúde", "acidente",
        "a quién llamo", "a quien llamo", "a dónde llamo", "adónde llamo", "telefono",
        "teléfono", "número", "numero", "contacto", "llamar", "llamo", "phone"
    ]

    complaint_terms = [
        "reclam", "queja", "complaint", "compens", "reembolso", "refund", "indemn"
    ]

    if any(term in q for term in critical_terms):
        return "critical_contact"
    if any(term in q for term in complaint_terms):
        return "complaint"
    return "general"


def detect_online_preference(question: str) -> bool:
    q = question.lower()
    online_terms = [
        "internet", "online", "enlace", "link", "web", "por internet",
        "llamada por internet", "llamar por internet", "solicitar llamada", "app", "sitio"
    ]
    return any(term in q for term in online_terms)


# ============================================================
# EXTRACCIÓN DE TELÉFONOS, ENLACES Y CONTACTOS
# ============================================================
def extract_phone_lines(text: str) -> List[str]:
    lines = []
    phone_pattern = r"\+?\d[\d\s\-()]{6,}\d"

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        low = line.lower()
        has_phone = re.search(phone_pattern, line) is not None
        has_contact_word = any(k in low for k in [
            "tel", "telefono", "teléfono", "phone", "whatsapp", "asistencia", "seguro",
            "europ assistance", "europassistance", "emergencia", "contacto", "24/7"
        ])

        if has_phone or has_contact_word:
            lines.append(line)

    dedup = []
    seen = set()
    for line in lines:
        key = normalize_for_match(line)
        if key not in seen:
            seen.add(key)
            dedup.append(line)
    return dedup[:15]


def extract_url_lines(text: str) -> List[str]:
    lines = []
    url_pattern = r"https?://\S+"

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        low = line.lower()
        has_url = re.search(url_pattern, line) is not None
        has_online_word = any(k in low for k in [
            "web", "enlace", "link", "internet", "online", "quickassistance", "eclaims",
            "solicitar una llamada", "llamada por internet"
        ])

        if has_url or has_online_word:
            lines.append(line)

    dedup = []
    seen = set()
    for line in lines:
        key = normalize_for_match(line)
        if key not in seen:
            seen.add(key)
            dedup.append(line)
    return dedup[:15]


def score_contact_relevance(question: str, chunk: Chunk) -> float:
    q = question.lower()
    text = f"{chunk.heading_path}\n{chunk.text}".lower()
    score = 0.0
    wants_online = detect_online_preference(question)

    heading_keywords = [
        "contactos críticos", "contactos", "emergencia", "seguro", "europ assistance",
        "actuación inmediata", "asistencia médica", "mi viaje", "accidente"
    ]
    for kw in heading_keywords:
        if kw in text:
            score += 0.12

    if any(k in q for k in ["accidente", "salud", "hospital", "médic", "medic", "seguro", "urgenc", "emergenc"]):
        for kw in ["europ assistance", "seguro", "asistencia médica", "ambul", "hospital", "urgencia", "emergencia"]:
            if kw in text:
                score += 0.08

    if any(k in q for k in ["llamo", "telefono", "teléfono", "contacto", "numero", "número"]):
        for kw in ["tel", "telefono", "teléfono", "contacto", "whatsapp", "24/7"]:
            if kw in text:
                score += 0.08

    if extract_phone_lines(chunk.text):
        score += 0.18

    if extract_url_lines(chunk.text):
        score += 0.10

    if wants_online:
        for kw in ["web", "quickassistance", "eclaims", "llamada por internet", "solicitar una llamada", "online", "internet"]:
            if kw in text:
                score += 0.16
        if extract_url_lines(chunk.text):
            score += 0.22

    return score


def build_priority_contact_block(question: str, retrieved: List[Tuple[Chunk, float]]) -> Dict[str, List[str]]:
    intent = detect_intent(question)
    if intent != "critical_contact":
        return {"contact_lines": [], "url_lines": [], "contact_chunks": []}

    rescored = []
    for chunk, base_score in retrieved:
        rescored.append((chunk, base_score + score_contact_relevance(question, chunk)))

    rescored.sort(key=lambda x: x[1], reverse=True)

    contact_lines = []
    url_lines = []
    contact_chunks = []
    seen_lines = set()
    seen_urls = set()
    seen_chunks = set()

    for chunk, _score in rescored[:CONTACTS_TOP_K]:
        phones = extract_phone_lines(chunk.text)
        urls = extract_url_lines(chunk.text)

        if (phones or urls) and chunk.chunk_id not in seen_chunks:
            seen_chunks.add(chunk.chunk_id)
            contact_chunks.append(chunk.heading_path)

        for line in phones:
            key = normalize_for_match(line)
            if key not in seen_lines:
                seen_lines.add(key)
                contact_lines.append(line)

        for line in urls:
            key = normalize_for_match(line)
            if key not in seen_urls:
                seen_urls.add(key)
                url_lines.append(line)

    return {
        "contact_lines": contact_lines[:12],
        "url_lines": url_lines[:10],
        "contact_chunks": contact_chunks[:6],
    }


# ============================================================
# LECTURA DE DOCUMENTOS
# ============================================================
def extract_text_from_txt(file_bytes: bytes) -> str:
    for encoding in ["utf-8", "latin-1", "cp1252"]:
        try:
            return file_bytes.decode(encoding)
        except Exception:
            continue
    return file_bytes.decode("utf-8", errors="ignore")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    if PdfReader is None:
        raise RuntimeError("Falta instalar PyPDF2 para procesar archivos PDF.")

    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for idx, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        pages.append(f"\n\n[Página {idx}]\n{page_text}")
    return "".join(pages)


def split_text_into_chunks(text: str, heading_path: str, source_name: str) -> List[Chunk]:
    text = normalize_whitespace(text)
    if not text:
        return []

    chunks: List[Chunk] = []
    start = 0
    text_len = len(text)

    while start < text_len:
        end = min(start + MAX_CHUNK_CHARS, text_len)
        candidate = text[start:end]

        if end < text_len:
            last_break = max(candidate.rfind("\n\n"), candidate.rfind(". "), candidate.rfind("; "))
            if last_break > int(MAX_CHUNK_CHARS * 0.55):
                end = start + last_break + 1
                candidate = text[start:end]

        chunk_text = normalize_whitespace(candidate)
        if chunk_text:
            chunk_id = hashlib.md5(f"{heading_path}|{start}|{chunk_text[:80]}".encode("utf-8")).hexdigest()
            chunks.append(Chunk(chunk_id=chunk_id, heading_path=heading_path, text=chunk_text, source_name=source_name))

        if end >= text_len:
            break
        start = max(end - CHUNK_OVERLAP_CHARS, start + 1)

    return chunks


def extract_sections_from_docx(file_bytes: bytes, source_name: str) -> List[Chunk]:
    if Document is None:
        raise RuntimeError("Falta instalar python-docx para procesar archivos Word.")

    doc = Document(io.BytesIO(file_bytes))
    sections: List[Tuple[str, List[str]]] = []
    heading_stack = [source_name]
    current_lines: List[str] = []
    last_heading_path = source_name

    for para in doc.paragraphs:
        text = normalize_whitespace(para.text)
        if not text:
            continue

        style_name = para.style.name if para.style else ""
        heading_match = re.match(r"Heading\s+(\d+)", style_name, re.IGNORECASE)
        titulo_match = re.match(r"T[íi]tulo\s+(\d+)", style_name, re.IGNORECASE)
        level = None

        if heading_match:
            level = int(heading_match.group(1))
        elif titulo_match:
            level = int(titulo_match.group(1))

        if level is not None:
            if current_lines:
                sections.append((last_heading_path, current_lines.copy()))
                current_lines = []

            heading_stack = heading_stack[:level]
            while len(heading_stack) < level:
                heading_stack.append("Sin título")
            if len(heading_stack) == level:
                heading_stack.append(text)
            else:
                heading_stack[level] = text
                heading_stack = heading_stack[: level + 1]

            last_heading_path = " > ".join([h for h in heading_stack if h])
        else:
            current_lines.append(text)

    if current_lines:
        sections.append((last_heading_path, current_lines.copy()))

    chunks: List[Chunk] = []
    for heading_path, lines in sections:
        full_text = normalize_whitespace("\n".join(lines))
        if full_text:
            chunks.extend(split_text_into_chunks(full_text, heading_path, source_name))

    if not chunks:
        raw_text = "\n".join([p.text for p in doc.paragraphs])
        chunks.extend(split_text_into_chunks(raw_text, source_name, source_name))

    return chunks


def extract_chunks_generic_text(file_bytes: bytes, source_name: str, extension: str) -> List[Chunk]:
    if extension == ".txt":
        raw_text = extract_text_from_txt(file_bytes)
    elif extension == ".pdf":
        raw_text = extract_text_from_pdf(file_bytes)
    else:
        raise ValueError("Formato no soportado en lectura genérica.")

    return split_text_into_chunks(normalize_whitespace(raw_text), source_name, source_name)


# ============================================================
# EMBEDDINGS Y RECUPERACIÓN
# ============================================================
def build_embeddings(client: OpenAI, chunks: List[Chunk], embedding_model: str) -> np.ndarray:
    texts = [f"Título: {c.heading_path}\n\nContenido: {c.text}" for c in chunks]
    vectors = []
    batch_size = 64

    for i in range(0, len(texts), batch_size):
        batch = texts[i:i + batch_size]
        resp = client.embeddings.create(model=embedding_model, input=batch)
        vectors.extend([item.embedding for item in resp.data])

    return np.array(vectors, dtype=np.float32)


def build_query_embedding(client: OpenAI, query: str, embedding_model: str) -> np.ndarray:
    resp = client.embeddings.create(model=embedding_model, input=query)
    return np.array(resp.data[0].embedding, dtype=np.float32)


def retrieve_top_chunks(
    client: OpenAI,
    question: str,
    chunks: List[Chunk],
    embeddings: np.ndarray,
    embedding_model: str,
    top_k: int = TOP_K,
) -> List[Tuple[Chunk, float]]:
    query_vec = build_query_embedding(client, question, embedding_model)
    scores = cosine_similarity_matrix(query_vec, embeddings)
    idxs = np.argsort(scores)[::-1][:top_k]
    return [(chunks[i], float(scores[i])) for i in idxs]


def boost_retrieval(question: str, retrieved: List[Tuple[Chunk, float]]) -> List[Tuple[Chunk, float]]:
    intent = detect_intent(question)
    wants_online = detect_online_preference(question)
    boosted = []

    for chunk, score in retrieved:
        bonus = 0.0
        text = f"{chunk.heading_path}\n{chunk.text}".lower()

        if intent == "critical_contact":
            if any(k in text for k in ["contactos críticos", "contactos", "europ assistance", "seguro", "asistencia médica", "actuación inmediata"]):
                bonus += 0.16
            if extract_phone_lines(chunk.text):
                bonus += 0.18
            if extract_url_lines(chunk.text):
                bonus += 0.12
            for kw in ["accidente", "salud", "hospital", "emergencia", "urgencia", "ambul", "seguro"]:
                if kw in text:
                    bonus += 0.07
            if wants_online:
                for kw in ["web", "quickassistance", "eclaims", "llamada por internet", "solicitar una llamada", "online", "internet"]:
                    if kw in text:
                        bonus += 0.20
                if extract_url_lines(chunk.text):
                    bonus += 0.20

        elif intent == "complaint":
            for kw in ["reclam", "compens", "reembolso", "cuestionario", "atención al cliente"]:
                if kw in text:
                    bonus += 0.06

        boosted.append((chunk, score + bonus))

    boosted.sort(key=lambda x: x[1], reverse=True)
    return boosted[:TOP_K]


# ============================================================
# RESPUESTA DEL MODELO
# ============================================================
def build_context_block(top_chunks: List[Tuple[Chunk, float]]) -> str:
    blocks = []
    for i, (chunk, score) in enumerate(top_chunks, start=1):
        blocks.append(
            f"[FUENTE {i}]\n"
            f"Ruta: {chunk.heading_path}\n"
            f"Documento: {chunk.source_name}\n"
            f"Relevancia: {score:.4f}\n"
            f"Contenido:\n{chunk.text}"
        )
    return "\n\n------------------------------\n\n".join(blocks)


def build_priority_notes(question: str, contact_block: Dict[str, List[str]]) -> str:
    intent = detect_intent(question)
    wants_online = detect_online_preference(question)
    notes = []

    if intent == "critical_contact":
        notes.append("La consulta es crítica y debe priorizar contactos, teléfonos, enlaces y actuación inmediata.")
        if wants_online:
            notes.append("El usuario pide una vía online o por internet. Prioriza enlaces web o canales online específicos antes que teléfonos alternativos no médicos.")
        if contact_block["url_lines"]:
            notes.append("Enlaces potencialmente relevantes detectados:")
            for line in contact_block["url_lines"][:6]:
                notes.append(f"- {line}")
        if contact_block["contact_lines"]:
            notes.append("Contactos potencialmente relevantes detectados:")
            for line in contact_block["contact_lines"][:8]:
                notes.append(f"- {line}")
        if not contact_block["contact_lines"] and not contact_block["url_lines"]:
            notes.append("No se detectaron contactos claros en los fragmentos priorizados.")

    elif intent == "complaint":
        notes.append("La consulta parece relacionada con reclamaciones, compensaciones o reembolsos. Prioriza reglas operativas y límites de actuación.")

    return "\n".join(notes) if notes else "Ninguna prioridad adicional detectada."


def ask_guidelines_assistant(
    client: OpenAI,
    question: str,
    top_chunks: List[Tuple[Chunk, float]],
    contact_block: Dict[str, List[str]],
    model_name: str,
    reasoning_effort: str,
) -> str:
    context_block = build_context_block(top_chunks)
    priority_notes = build_priority_notes(question, contact_block)
    wants_online = detect_online_preference(question)

    instructions = (
        "Eres un asistente interno para guías de una empresa de viajes. "
        "Tu única fuente válida es el contexto documental proporcionado. "
        "No inventes políticas, procesos, teléfonos, enlaces ni excepciones. "
        "Debes responder en el mismo idioma en que el usuario formule la consulta. "
        "La respuesta debe ser breve, precisa, operativa y útil para actuar de inmediato. "
        "Analiza bien el contexto antes de responder y prioriza el dato más accionable. "
        "Cuando la consulta trate sobre accidente, salud, urgencia, hospital, seguro o 'a quién llamo', "
        "debes priorizar el contacto aplicable y colocarlo en la primera línea si está presente en el contexto. "
        "Si el usuario pregunta por una vía online, por internet, web o enlace, debes priorizar el enlace específico del caso si aparece en el contexto. "
        "No sustituyas un enlace médico específico por teléfonos o WhatsApps generales si el documento ofrece una vía online médica directa. "
        "Si existen varios contactos, muestra primero el más directamente relacionado con el caso. "
        "No redactes párrafos extensos. Usa el siguiente formato, siempre que el contexto lo permita: 'Contacto inmediato' y 'Qué hacer ahora'. "
        "No añadas una sección final de fuentes o base documental en la respuesta visible al usuario. "
        "Si no hay contacto identificable, dilo expresamente y da solo los pasos indispensables."
    )

    focus_rule = (
        "7. Si el usuario pide una vía por internet, online, web o enlace, prioriza el enlace web específico y no lo omitas si aparece en el contexto.\n"
        if wants_online else
        "7. Si el contexto contiene un dato directo y evidente, priorízalo sobre explicaciones generales.\n"
    )

    user_input = (
        f"CONSULTA DEL GUÍA:\n{question}\n\n"
        f"PRIORIDADES OPERATIVAS:\n{priority_notes}\n\n"
        f"CONTEXTO DOCUMENTAL DISPONIBLE:\n{context_block}\n\n"
        "Reglas de salida obligatorias:\n"
        "1. Respuesta preferente entre 70 y 150 palabras.\n"
        "2. Si hay un teléfono, enlace o contacto útil, ponlo en la primera línea bajo el rótulo 'Contacto inmediato'.\n"
        "3. Después incluye 'Qué hacer ahora' con 2 a 4 viñetas como máximo.\n"
        "4. No añadas una sección final de fuentes o base documental en la respuesta visible al usuario.\n"
        "5. No copies párrafos largos del contexto.\n"
        "6. No añadas recomendaciones externas no contenidas en el documento.\n"
        f"{focus_rule}"
    )

    response = client.responses.create(
        model=model_name,
        reasoning={"effort": reasoning_effort},
        instructions=instructions,
        input=user_input,
    )

    return response.output_text.strip()


# ============================================================
# ESTADO Y PROCESAMIENTO
# ============================================================
def initialize_cache_state() -> None:
    if "doc_cache" not in st.session_state:
        st.session_state.doc_cache = {}
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "last_file_hash" not in st.session_state:
        st.session_state.last_file_hash = None


def reset_chat_if_document_changed(current_hash: str) -> None:
    if st.session_state.last_file_hash != current_hash:
        st.session_state.messages = []
        st.session_state.last_file_hash = current_hash


def process_document(client: OpenAI, source_file, embedding_model: str):
    file_bytes = source_file.read()
    file_hash = file_sha256(file_bytes)
    source_name = source_file.name
    extension = os.path.splitext(source_name)[1].lower()
    cache_key = f"{file_hash}|{embedding_model}"

    if cache_key in st.session_state.doc_cache:
        return st.session_state.doc_cache[cache_key]

    if extension == ".docx":
        chunks = extract_sections_from_docx(file_bytes, source_name)
    elif extension in [".pdf", ".txt"]:
        chunks = extract_chunks_generic_text(file_bytes, source_name, extension)
    else:
        raise ValueError("Formato no soportado. Usa .docx, .pdf o .txt")

    if not chunks:
        raise ValueError("No se pudo extraer contenido útil del documento.")

    embeddings = build_embeddings(client, chunks, embedding_model)

    payload = {
        "file_hash": file_hash,
        "source_name": source_name,
        "chunks": chunks,
        "embeddings": embeddings,
        "processed_at": time.time(),
        "extension": extension,
    }
    st.session_state.doc_cache[cache_key] = payload
    return payload


# ============================================================
# INTERFAZ
# ============================================================
def render_css() -> None:
    st.markdown(
        """
        <style>
            [data-testid="stSidebar"] {display: none;}
            .block-container {padding-top: 1.4rem; padding-bottom: 5rem; max-width: 1000px;}
        </style>
        """,
        unsafe_allow_html=True,
    )


def render_header():
    st.title(APP_TITLE)
    st.caption(APP_SUBTITLE)


def render_history():
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])


def main():
    initialize_cache_state()
    render_css()
    render_header()

    api_key = safe_get_api_key()
    if not api_key:
        st.error("No se encontró OPENAI_API_KEY en secrets.toml.")
        st.stop()

    if not os.path.exists(REPO_DOC_PATH):
        st.error(f"No se encontró el documento en el repositorio: {REPO_DOC_PATH}")
        st.stop()

    client = OpenAI(api_key=api_key)
    model_name = DEFAULT_MODEL
    reasoning_effort = DEFAULT_REASONING_EFFORT
    embedding_model = DEFAULT_EMBEDDING_MODEL
    repo_file = RepoFile(REPO_DOC_PATH)

    try:
        with st.spinner("Procesando documento..."):
            doc_data = process_document(client, repo_file, embedding_model)
    except Exception as e:
        st.error(f"No se pudo procesar el documento: {e}")
        st.stop()

    reset_chat_if_document_changed(doc_data["file_hash"])
    render_history()

    question = st.chat_input("Escribe la consulta del guía...")

    if question:
        st.session_state.messages.append({"role": "user", "content": question})

        with st.chat_message("user"):
            st.markdown(question)

        with st.chat_message("assistant"):
            try:
                with st.spinner("Consultando la base documental..."):
                    retrieved = retrieve_top_chunks(
                        client=client,
                        question=question,
                        chunks=doc_data["chunks"],
                        embeddings=doc_data["embeddings"],
                        embedding_model=embedding_model,
                        top_k=TOP_K,
                    )
                    top_chunks = boost_retrieval(question, retrieved)
                    contact_block = build_priority_contact_block(question, top_chunks)
                    answer = ask_guidelines_assistant(
                        client=client,
                        question=question,
                        top_chunks=top_chunks,
                        contact_block=contact_block,
                        model_name=model_name,
                        reasoning_effort=reasoning_effort,
                    )

                st.markdown(answer)
                st.session_state.messages.append({"role": "assistant", "content": answer})

            except Exception as e:
                error_msg = f"Se produjo un error al generar la respuesta: {e}"
                st.error(error_msg)
                st.session_state.messages.append({"role": "assistant", "content": error_msg})


if __name__ == "__main__":
    main()
